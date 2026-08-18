#!/usr/bin/env python3
"""
Pre-publish self-check for content authored against IS 5568.

    python preflight.py <file-or-directory> [--json]

Deliberately conservative about what it claims. It reports three things
separately:

    FAIL      a defect it can demonstrate
    CHECK     something it cannot decide, that a human must look at
    OK        a check that ran and found nothing

A clean run is **not** compliance. Roughly two-thirds of the criteria need human
judgement, and this script says so rather than implying a pass. For a full audit
against all 60 checks, use the is5568-auditor tool.

Only the standard library is required for HTML and text. Office and PDF checks
need the optional dependencies noted at runtime.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import zipfile
from dataclasses import dataclass, field
from html.parser import HTMLParser
from typing import Any

HEBREW_RE = re.compile(r"[֐-׿]")
LATIN_RE = re.compile(r"[A-Za-z]")
GENERIC_LINK = {
    "לחץ כאן", "לחצו כאן", "כאן", "קרא עוד", "קראו עוד", "למידע נוסף",
    "לפרטים", "עוד", "המשך", "click here", "read more", "more", "here", "link",
}


@dataclass
class Report:
    target: str
    kind: str
    failures: list[str] = field(default_factory=list)
    checks: list[str] = field(default_factory=list)
    passed: list[str] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)

    def fail(self, msg: str) -> None:
        self.failures.append(msg)

    def check(self, msg: str) -> None:
        self.checks.append(msg)

    def ok(self, msg: str) -> None:
        self.passed.append(msg)


# ── HTML ────────────────────────────────────────────────────────────────────

class HtmlScan(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.html_attrs: dict[str, str] = {}
        self.title = ""
        self._in_title = False
        self.headings: list[tuple[int, str]] = []
        self._heading_level: int | None = None
        self._heading_text = ""
        self.images: list[dict[str, Any]] = []
        self.inputs: list[dict[str, Any]] = []
        self.labels_for: set[str] = set()
        self.links: list[dict[str, Any]] = []
        self._link_attrs: dict[str, str] | None = None
        self._link_text = ""
        self.ids: list[str] = []
        self.landmarks: set[str] = set()
        self.tables = 0
        self.th = 0
        self.captions = 0
        self.lang_spans = 0
        self.iframes: list[dict[str, str]] = []
        self.presentational: list[str] = []
        self.viewport = ""
        self.buttons_without_text = 0
        self._button_attrs: dict[str, str] | None = None
        self._button_text = ""

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        a = {k.lower(): (v or "") for k, v in attrs}
        if "id" in a:
            self.ids.append(a["id"])
        if "lang" in a and tag != "html":
            self.lang_spans += 1

        if tag == "html":
            self.html_attrs = a
        elif tag == "title":
            self._in_title = True
        elif tag == "meta" and a.get("name", "").lower() == "viewport":
            self.viewport = a.get("content", "")
        elif re.fullmatch(r"h[1-6]", tag):
            self._heading_level = int(tag[1])
            self._heading_text = ""
        elif tag in ("img", "area", "input") and (tag != "input" or a.get("type") == "image"):
            self.images.append(a)
        elif tag == "input":
            self.inputs.append(a)
        elif tag in ("select", "textarea"):
            self.inputs.append(a)
        elif tag == "label" and "for" in a:
            self.labels_for.add(a["for"])
        elif tag == "a":
            self._link_attrs = a
            self._link_text = ""
        elif tag == "button":
            self._button_attrs = a
            self._button_text = ""
        elif tag in ("header", "nav", "main", "footer", "aside"):
            self.landmarks.add(tag)
        elif tag == "table":
            self.tables += 1
        elif tag == "th":
            self.th += 1
        elif tag == "caption":
            self.captions += 1
        elif tag == "iframe":
            self.iframes.append(a)
        elif tag in ("font", "center", "marquee", "blink", "big", "strike", "tt"):
            self.presentational.append(tag)

        if a.get("role") in ("banner", "navigation", "main", "contentinfo"):
            self.landmarks.add(a["role"])

    def handle_endtag(self, tag: str) -> None:
        if tag == "title":
            self._in_title = False
        elif re.fullmatch(r"h[1-6]", tag) and self._heading_level is not None:
            self.headings.append((self._heading_level, self._heading_text.strip()))
            self._heading_level = None
        elif tag == "a" and self._link_attrs is not None:
            self.links.append({**self._link_attrs, "_text": self._link_text.strip()})
            self._link_attrs = None
        elif tag == "button" and self._button_attrs is not None:
            has_name = bool(self._button_text.strip()) or bool(
                self._button_attrs.get("aria-label") or self._button_attrs.get("aria-labelledby")
            )
            if not has_name:
                self.buttons_without_text += 1
            self._button_attrs = None

    def handle_data(self, data: str) -> None:
        if self._in_title:
            self.title += data
        if self._heading_level is not None:
            self._heading_text += data
        if self._link_attrs is not None:
            self._link_text += data
        if self._button_attrs is not None:
            self._button_text += data


def check_html(path: str, report: Report) -> None:
    with open(path, "rb") as fh:
        raw = fh.read()
    text = raw.decode("utf-8", "replace")

    scan = HtmlScan()
    try:
        scan.feed(text)
    except Exception as exc:  # noqa: BLE001
        report.fail(f"לא ניתן היה לנתח את ה-HTML: {exc}")
        return

    body_text = re.sub(r"<[^>]+>", " ", text)
    is_rtl_content = len(HEBREW_RE.findall(body_text)) > len(LATIN_RE.findall(body_text))

    # 3.1.1 / Israeli addition — language and direction
    lang = scan.html_attrs.get("lang", "")
    if not lang:
        report.fail("3.1.1 — לא הוגדר מאפיין lang בתגית html")
    else:
        report.ok(f"3.1.1 — שפת המסמך מוצהרת ({lang})")
    if is_rtl_content and scan.html_attrs.get("dir") != "rtl":
        report.fail('תוספת ישראלית — תוכן בעברית ללא dir="rtl" בתגית html')
    elif is_rtl_content:
        report.ok('כיוון הכתיבה מוצהר (dir="rtl")')

    # 2.4.2 — page title
    if not scan.title.strip():
        report.fail("2.4.2 — לעמוד אין כותרת (title)")
    else:
        report.ok(f"2.4.2 — קיימת כותרת: {scan.title.strip()[:60]!r}")
        report.check("2.4.2 — ודאו שהכותרת מתארת את העמוד הספציפי וייחודית באתר")

    # 1.3.1 / 2.4.10 — heading structure
    levels = [lvl for lvl, _ in scan.headings]
    if not levels:
        report.fail("1.3.1 — לא נמצאו כותרות (h1-h6) בעמוד")
    else:
        if levels.count(1) == 0:
            report.fail("2.4.10 — לא נמצאה כותרת ראשית (h1)")
        elif levels.count(1) > 1:
            report.fail(f"2.4.10 — נמצאו {levels.count(1)} כותרות h1; נדרשת אחת בלבד")
        else:
            report.ok("2.4.10 — קיימת כותרת ראשית אחת")
        prev = 0
        for lvl, txt in scan.headings:
            if prev and lvl > prev + 1:
                report.fail(f"1.3.1 — דילוג ברמות הכותרות: h{lvl} ({txt[:40]!r}) אחרי h{prev}")
            prev = lvl
        for lvl, txt in scan.headings:
            if not txt:
                report.fail(f"1.3.1 — כותרת h{lvl} ריקה")

    # 1.1.1 — images
    missing_alt = [i for i in scan.images if "alt" not in i and not i.get("aria-label")]
    filename_alt = [i for i in scan.images if re.search(r"\.(jpe?g|png|gif|svg|webp)$", i.get("alt", ""), re.I)]
    if scan.images:
        if missing_alt:
            report.fail(f"1.1.1 — {len(missing_alt)} תמונות ללא מאפיין alt")
        if filename_alt:
            report.fail(f"1.1.1 — {len(filename_alt)} תמונות שה-alt שלהן הוא שם קובץ")
        if not missing_alt and not filename_alt:
            report.ok(f"1.1.1 — לכל {len(scan.images)} התמונות יש טקסט חלופי")
        report.check("1.1.1 — ודאו ידנית שכל טקסט חלופי מתאר בפועל את מה שהתמונה מעבירה")
    else:
        report.ok("1.1.1 — לא נמצאו תמונות")

    # 3.3.2 — form labels
    labelled_targets = scan.labels_for
    unlabelled = [
        i for i in scan.inputs
        if i.get("type") not in ("hidden", "submit", "button", "reset")
        and not i.get("aria-label")
        and not i.get("aria-labelledby")
        and i.get("id", "\0") not in labelled_targets
    ]
    if scan.inputs:
        if unlabelled:
            report.fail(f"3.3.2 — {len(unlabelled)} שדות טופס ללא תווית משויכת")
            placeholder_only = [i for i in unlabelled if i.get("placeholder")]
            if placeholder_only:
                report.fail(
                    f"3.3.2 — {len(placeholder_only)} מהם משתמשים ב-placeholder בלבד; "
                    "ה-placeholder נעלם בהקלדה ואינו תחליף ל-label"
                )
        else:
            report.ok(f"3.3.2 — לכל {len(scan.inputs)} שדות הטופס יש תווית")

    # 2.4.4 — link purpose
    generic = [l for l in scan.links if l.get("_text", "").strip().lower() in GENERIC_LINK]
    empty = [l for l in scan.links if not l.get("_text", "").strip() and not l.get("aria-label")]
    if empty:
        report.fail(f"2.4.4 / 4.1.2 — {len(empty)} קישורים ללא טקסט או שם נגיש")
    if generic:
        report.check(
            f"2.4.4 — {len(generic)} קישורים עם טקסט גנרי ({generic[0].get('_text')!r}). "
            "מותר רק אם המשפט שסביבם מבהיר את היעד"
        )
    if scan.links and not empty and not generic:
        report.ok(f"2.4.4 — {len(scan.links)} קישורים, כולם עם טקסט תיאורי")

    # 4.1.2 — names for controls and frames
    if scan.buttons_without_text:
        report.fail(f"4.1.2 — {scan.buttons_without_text} כפתורים ללא שם נגיש")
    untitled_frames = [f for f in scan.iframes if not f.get("title")]
    if untitled_frames:
        report.fail(f"4.1.2 — {len(untitled_frames)} מסגרות iframe ללא title")

    # 4.1.1 — duplicate ids
    dupes = {i for i in scan.ids if scan.ids.count(i) > 1}
    if dupes:
        report.fail(f"4.1.1 — מזהים כפולים: {', '.join(sorted(dupes)[:5])}")
    else:
        report.ok("4.1.1 — כל המזהים ייחודיים")

    # 1.3.1 — content/presentation separation
    if scan.presentational:
        report.fail(
            f"1.3.1 — תגיות עיצוביות בקוד: {', '.join(sorted(set(scan.presentational)))}"
        )

    # 1.4.4 — zoom
    if re.search(r"user-scalable\s*=\s*(no|0)|maximum-scale\s*=\s*1", scan.viewport, re.I):
        report.fail("1.4.4 — תגית viewport חוסמת שינוי מרחק מתצוגה")
    elif scan.viewport:
        report.ok("1.4.4 — שינוי מרחק מתצוגה אינו חסום")

    # 2.4.1 — bypass mechanism
    skip = [l for l in scan.links if l.get("href", "").startswith("#") and re.search(r"דלג|skip", l.get("_text", ""), re.I)]
    if skip:
        report.ok("2.4.1 — נמצא קישור דילוג")
    elif {"main"} & scan.landmarks:
        report.ok("2.4.1 — קיים אזור main המשמש כמנגנון עקיפה")
    else:
        report.fail("2.4.1 — לא נמצא קישור דילוג ולא אזור main")

    # 1.3.1 — tables
    if scan.tables:
        if scan.th == 0:
            report.fail(f"1.3.1 — {scan.tables} טבלאות ללא תאי כותרת (th)")
        if scan.captions == 0:
            report.check(f"1.3.1 — {scan.tables} טבלאות ללא caption")

    # Israeli addition — accessibility statement
    if re.search(r"הצהרת נגישות", text):
        report.ok("תוספת ישראלית — נמצא קישור/אזכור להצהרת נגישות")
    else:
        report.fail("תוספת ישראלית — לא נמצא קישור להצהרת נגישות")

    # Israeli addition — English accessible names on a Hebrew page
    if is_rtl_content:
        english_labels = [
            m for m in re.findall(r'aria-label="([^"]+)"', text)
            if LATIN_RE.search(m) and not HEBREW_RE.search(m) and len(m.split()) > 1
        ]
        if english_labels:
            report.fail(
                f"תוספת ישראלית — שמות נגישים באנגלית בעמוד בעברית: {english_labels[:3]}"
            )

    # Things this script structurally cannot decide.
    report.check("1.4.3 — ניגודיות נבדקת רק בדפדפן; הריצו את is5568-auditor")
    report.check("2.1.1 / 2.1.2 / 2.4.3 / 2.4.7 — ניווט ומיקוד מקלדת דורשים בדיקה בדפדפן")
    report.check("1.3.2 — סדר הקריאה מול הסדר החזותי דורש רינדור")
    report.skipped.append("קריטריונים התלויים ברינדור נבדקים על ידי is5568-auditor בלבד")


# ── Office and PDF ──────────────────────────────────────────────────────────

def check_docx(path: str, report: Report) -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        report.skipped.append("python-docx לא מותקן — דילוג על בדיקת ה-Word")
        return
    import docx

    document = docx.Document(path)
    core = document.core_properties

    if (core.title or "").strip():
        report.ok(f"2.4.2 — למסמך יש כותרת: {core.title!r}")
    else:
        base = os.path.splitext(os.path.basename(path))[0]
        if re.fullmatch(r"(doc|document|scan|final|copy|new|untitled|מסמך)[\s_-]*\d*", base, re.I):
            report.fail(f"2.4.2 — אין כותרת למסמך ושם הקובץ {base!r} אינו בעל משמעות")
        else:
            report.check("2.4.2 — אין מאפיין Title; שם הקובץ משמש כשם המסמך")

    styled = [p for p in document.paragraphs
              if p.style and (p.style.name.startswith("Heading") or p.style.name.startswith("כותרת"))]
    fake = []
    for p in document.paragraphs:
        txt = (p.text or "").strip()
        if not txt or len(txt) > 120:
            continue
        if p.style and (p.style.name.startswith("Heading") or p.style.name.startswith("כותרת")):
            continue
        runs = [r for r in p.runs if (r.text or "").strip()]
        if not runs:
            continue
        first = runs[0]
        size = first.font.size.pt if first.font.size else None
        if first.font.bold or (size and size >= 14):
            fake.append(txt)

    if styled:
        report.ok(f"1.3.1 — {len(styled)} כותרות מסומנות בסגנון כותרת")
    else:
        report.fail("1.3.1 — לא נמצאו כותרות בסגנון כותרת סמנטי במסמך")
    if fake:
        report.fail(
            f"1.3.1 — {len(fake)} פסקאות מעוצבות ככותרת ללא סגנון כותרת "
            f"(לדוגמה: {fake[0][:50]!r})"
        )

    images = docx_images(path)
    no_alt = [i for i in images if not i["alt"] and not i["decorative"]]
    if images:
        if no_alt:
            report.fail(f"1.1.1 — {len(no_alt)} תמונות במסמך ללא טקסט חלופי")
        else:
            report.ok(f"1.1.1 — לכל {len(images)} התמונות טקסט חלופי או סימון דקורטיבי")

    for i, table in enumerate(document.tables, start=1):
        try:
            marked = bool(table.rows[0]._tr.findall(  # noqa: SLF001
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader"
            ))
        except Exception:  # noqa: BLE001
            marked = False
        if not marked:
            report.fail(f"1.3.1 — טבלה {i} ללא שורת כותרת מסומנת (Repeat Header Rows)")

    report.check("1.4.3 — ניגודיות במסמך: ספי חלק 2 הם 14 נק' מודגש / 18 נק' רגיל")
    report.check("1.3.2 — ודאו סדר קריאה: תיבות טקסט צפות נקראות מחוץ לרצף")


def docx_images(path: str) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    try:
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if not (name.startswith("word/") and name.endswith(".xml")):
                    continue
                xml = zf.read(name).decode("utf-8", "ignore")
                for match in re.finditer(r"<wp:docPr\b([^>]*)/?>", xml):
                    attrs = match.group(1)
                    descr = re.search(r'descr="([^"]*)"', attrs)
                    out.append({
                        "alt": descr.group(1) if descr else "",
                        "decorative": "decorative" in attrs.lower(),
                    })
    except Exception:  # noqa: BLE001
        pass
    return out


def check_pdf(path: str, report: Report) -> None:
    try:
        import pikepdf
    except ImportError:
        report.skipped.append("pikepdf לא מותקן — דילוג על בדיקת ה-PDF")
        return

    with pikepdf.open(path) as pdf:
        root = pdf.Root
        marked = bool(root.get("/MarkInfo", {}).get("/Marked", False))
        if marked and root.get("/StructTreeRoot") is not None:
            report.ok("1.3.1 — המסמך מתויג ויש בו עץ תגיות")
        elif marked:
            report.fail("1.3.1 — המסמך מסומן כמתויג אך אין בו עץ תגיות (StructTreeRoot)")
        else:
            report.fail(
                "1.3.1 / 1.3.2 — המסמך אינו מתויג. אין מבנה, אין סדר קריאה, ואין היכן לשאת טקסט חלופי. "
                "יש לייצא מחדש עם 'Document structure tags for accessibility'"
            )

        if root.get("/Lang"):
            report.ok(f"3.1.1 — שפת המסמך מוצהרת ({root.get('/Lang')})")
        else:
            report.fail("3.1.1 — לא הוגדרה שפת המסמך (/Lang)")

        title = None
        with pdf.open_metadata() as meta:
            title = meta.get("dc:title")
        if not title and pdf.docinfo:
            title = pdf.docinfo.get("/Title")
        if title:
            report.ok(f"2.4.2 — למסמך יש כותרת: {str(title)[:60]!r}")
        else:
            report.fail("2.4.2 — למסמך אין כותרת")

        prefs = root.get("/ViewerPreferences")
        if prefs is not None and prefs.get("/DisplayDocTitle"):
            report.ok("2.4.2 — DisplayDocTitle מופעל; הקורא יכריז את הכותרת")
        else:
            report.fail(
                "2.4.2 — DisplayDocTitle כבוי; קוראי PDF יציגו את שם הקובץ במקום את הכותרת"
            )

    try:
        import fitz

        doc = fitz.open(path)
        total = sum(len((doc[i].get_text("text") or "").strip()) for i in range(doc.page_count))
        pages = doc.page_count
        doc.close()
        if pages and total < 40:
            report.fail(
                "1.4.5 — לא נמצא טקסט הניתן לחילוץ: המסמך סרוק. "
                "חלק 2 אוסר במפורש שימוש בקבצים סרוקים"
            )
        else:
            report.ok(f"1.4.5 — נמצא טקסט הניתן לחילוץ ({total} תווים ב-{pages} עמודים)")
    except ImportError:
        report.skipped.append("pymupdf לא מותקן — לא נבדק אם המסמך סרוק")

    report.check("1.4.3 — ניגודיות: ספי חלק 2 הם 14 נק' מודגש / 18 נק' רגיל")
    report.check("1.3.2 — ודאו את סדר הקריאה בלוח התגיות של הכלי שבו נוצר המסמך")


CHECKERS = {
    ".html": check_html,
    ".htm": check_html,
    ".docx": check_docx,
    ".pdf": check_pdf,
}


def run(path: str) -> Report:
    ext = os.path.splitext(path)[1].lower()
    report = Report(target=path, kind=ext.lstrip(".") or "unknown")
    checker = CHECKERS.get(ext)
    if checker is None:
        report.skipped.append(f"אין בדיקת preflight לסיומת {ext!r}")
        return report
    try:
        checker(path, report)
    except Exception as exc:  # noqa: BLE001
        report.fail(f"הבדיקה נכשלה טכנית: {exc}")
    return report


def collect(target: str) -> list[str]:
    if os.path.isfile(target):
        return [target]
    out: list[str] = []
    for root, _dirs, files in os.walk(target):
        for name in files:
            if os.path.splitext(name)[1].lower() in CHECKERS:
                out.append(os.path.join(root, name))
    return sorted(out)


def main() -> int:
    parser = argparse.ArgumentParser(description="IS 5568 pre-publish self-check.")
    parser.add_argument("target", help="file or directory to check")
    parser.add_argument("--json", action="store_true", help="machine-readable output")
    args = parser.parse_args()

    if not os.path.exists(args.target):
        print(f"not found: {args.target}", file=sys.stderr)
        return 2

    paths = collect(args.target)
    if not paths:
        print(f"no checkable files under {args.target}", file=sys.stderr)
        return 2

    reports = [run(p) for p in paths]

    if args.json:
        json.dump([r.__dict__ for r in reports], sys.stdout, ensure_ascii=False, indent=2)
        print()
        return 1 if any(r.failures for r in reports) else 0

    total_fail = 0
    for r in reports:
        print(f"\n{'=' * 70}\n{r.target}  [{r.kind}]\n{'=' * 70}")
        for msg in r.failures:
            print(f"  FAIL   {msg}")
        for msg in r.checks:
            print(f"  CHECK  {msg}")
        for msg in r.passed:
            print(f"  OK     {msg}")
        for msg in r.skipped:
            print(f"  ---    {msg}")
        total_fail += len(r.failures)

    print(f"\n{'=' * 70}")
    print(f"{len(reports)} files · {total_fail} failures · "
          f"{sum(len(r.checks) for r in reports)} items needing human review")
    print(
        "\nA clean run is not compliance. Roughly two-thirds of the criteria need\n"
        "human judgement — run the is5568-auditor tool for the full check, and test\n"
        "with a screen reader before publishing."
    )
    return 1 if total_fail else 0


if __name__ == "__main__":
    sys.exit(main())
